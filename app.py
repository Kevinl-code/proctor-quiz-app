from flask import Flask, render_template, request, redirect, flash, jsonify, session
from pymongo import MongoClient
import re
import uuid
from datetime import datetime, timedelta
import pandas as pd
import docx
import pdfplumber
import os
import io
import qrcode
from io import BytesIO
from flask import send_file
from flask import send_from_directory
from twilio.twiml.messaging_response import MessagingResponse
from requests.auth import HTTPBasicAuth
import requests


app = Flask(__name__)
app.secret_key = 'supersecretkey'

# ================= DATABASE =================
client = MongoClient("mongodb://Kevin2003:%40Kevin2003@ac-gjdvsbl-shard-00-00.gpbpget.mongodb.net:27017,ac-gjdvsbl-shard-00-01.gpbpget.mongodb.net:27017,ac-gjdvsbl-shard-00-02.gpbpget.mongodb.net:27017/?ssl=true&replicaSet=atlas-dc5j3m-shard-0&authSource=admin&appName=proctor")

db = client['proctor']

users_collection = db['users']
quiz = db["quizzes"]
questions = db["questions"]
activity = db["student_activity"]
scores = db["scores"]
submissions = db["submissions"]

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ================= EMAIL ROLE =================
teacher_pattern = re.compile(r'^[a-z0-9]+@bhc\.professor\.com$')
student_pattern = re.compile(r'^[a-z0-9]+@bhc\.student\.com$')

# ================= HOME =================
@app.route('/')
def spinner():
    return render_template('spinner.html')

# ================= LOGIN =================
@app.route("/login", methods=["GET","POST"])
def login():

    if request.method == "POST":

        email = request.form["email"]
        password = request.form["password"]

        user = users_collection.find_one({"email":email,"password":password})

        if user:

            session["user"] = email  
            session["name"] = user["name"]  # ✅ FIX SESSION

            if teacher_pattern.match(email):
                return redirect("/admin")
            elif student_pattern.match(email):
                return redirect("/student")

        flash("Invalid credentials")
        return redirect("/login")

    return render_template("login.html")

# ================= LOGOUT =================
@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")
# ================= SIGNUP =================
@app.route("/signup",methods=["GET","POST"])
def signup():

    if request.method == "POST":

        data={
            "name":request.form["name"],
            "email":request.form["email"],
            "password":request.form["password"],
            "role":request.form["role"]
        }

        users_collection.insert_one(data)

        flash("Account Created Successfully")
        return redirect("/login")

    return render_template("signup.html")


@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static/images'),
                               'favicon.png', mimetype='image/vnd.microsoft.icon')

# ================= ADMIN =================
@app.route("/admin")
def admin_dashboard():

    if "user" not in session:
        return redirect("/login")

    return render_template("admin_dashboard.html")

# ================= STUDENT =================
@app.route("/student")
def student_dashboard():

    if "user" not in session:
        return redirect("/login")

    return render_template("student_dashboard.html")

# ================= CREATE QUIZ =================
@app.route("/create_quiz",methods=["POST"])
def create_quiz():

    data=request.json

    quiz_id=str(uuid.uuid4())[:8]

    start_time = datetime.fromisoformat(data["start"])
    duration = int(data["duration"])

    end_time = start_time + timedelta(minutes=duration)

    quiz_doc={
        "quiz_id":quiz_id,
        "title":data["title"],
        "start_time":start_time.isoformat(),
        "end_time":end_time.isoformat(),
        "duration":duration,
        "created_at":datetime.now()
    }

    quiz.insert_one(quiz_doc)

    # save questions
    for q in data.get("questions",[]):

        questions.insert_one({
            "quiz_id":quiz_id,
            "question":q["question"],
            "options":q["options"],
            "answer":q["answer"]
        })

    return jsonify({"msg": f'{data["title"]} Created Successfully',"quiz_id": quiz_id})
@app.route("/generate_qr/<quiz_id>")
def generate_qr(quiz_id):

    url = request.host_url + "quiz_info/" + quiz_id

    qr = qrcode.make(url)

    img_io = io.BytesIO()
    qr.save(img_io, 'PNG')
    img_io.seek(0)

    return send_file(img_io, mimetype='image/png')

@app.route("/quiz_info/<quiz_id>")
def quiz_info(quiz_id):

    # 🔐 NOT LOGGED IN → LOGIN
    if "user" not in session:
        return redirect("/login")

    # 👨‍🏫 ADMIN BLOCK
    email = session.get("user")
    if teacher_pattern.match(email):
        return redirect("/admin")

    # 👨‍🎓 STUDENT → SHOW QUIZ INFO PAGE
    q = quiz.find_one({"quiz_id": quiz_id}, {"_id":0})

    return render_template("quiz_info.html", quiz=q)
def attend_quiz(quiz_id):

    # 🔐 NOT LOGGED
    if "user" not in session:
        return redirect("/login")

    email = session.get("user")

    # 👨‍🏫 BLOCK ADMIN
    if teacher_pattern.match(email):
        return redirect("/admin")

    return render_template("student_quiz.html")

@app.route("/join/<quiz_id>")
def join_quiz(quiz_id):

    user = session.get("user")

    if not user:
        return redirect("/login")

    if teacher_pattern.match(user):
        return redirect("/admin")

    return redirect(f"/quiz/{quiz_id}")
    
# ================= FILE UPLOAD =================
@app.route("/upload_questions", methods=["POST"])
def upload_questions():

    file = request.files.get("file")

    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    filename = file.filename.lower()
    parsed = []

    # ================= CSV =================
    if filename.endswith(".csv"):
        df = pd.read_csv(file)

        for _, r in df.iterrows():
            parsed.append({
                "question": str(r["question"]),
                "options": [r["A"], r["B"], r["C"], r["D"]],
                "answer": str(r["answer"]).strip().upper()
            })

    # ================= TXT =================
    elif filename.endswith(".txt"):
        lines = file.read().decode("utf-8").split("\n")

        for line in lines:
            parts = line.strip().split("|")

            if len(parts) == 6:
                parsed.append({
                    "question": parts[0],
                    "options": parts[1:5],
                    "answer": parts[5].strip().upper()
                })

    # ================= DOCX =================
    elif filename.endswith(".docx"):
        doc = docx.Document(file)

        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        parsed.extend(parse_block_questions(lines))

    # ================= PDF =================
    elif filename.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:

            lines = []

            for p in pdf.pages:
                text = p.extract_text()
                if text:
                    lines.extend(text.split("\n"))

        parsed.extend(parse_block_questions(lines))

    else:
        return jsonify({"error": "Unsupported file"}), 400

    return jsonify(parsed)
# ================= PARSER =================
def parse_block_questions(lines):

    result = []
    current = None

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # QUESTION
        if "?" in line:
            if current and len(current["options"]) == 4 and current["answer"]:
                result.append(current)

            current = {
                "question": line,
                "options": [],
                "answer": ""
            }

        # OPTIONS
        elif line.startswith(("A.", "B.", "C.", "D.")):
            if current:
                current["options"].append(line[2:].strip())

        # ANSWER
        elif "answer" in line.lower():
            if current:
                ans = line.split(":")[-1].strip().upper()
                if ans in ["A","B","C","D"]:
                    current["answer"] = ans

    # LAST QUESTION
    if current and len(current["options"]) == 4 and current["answer"]:
        result.append(current)

    return result
# ================= GET QUIZZES =================
@app.route("/get_quizzes")
def get_quizzes():

    quizzes = list(quiz.find({}, {"_id":0}))

    student_id = session.get("user")

    now = datetime.now()

    for q in quizzes:

        start = datetime.fromisoformat(q["start_time"])
        end = datetime.fromisoformat(q["end_time"])

        attempt = submissions.find_one({
            "quiz_id": q["quiz_id"],
            "student_id": student_id
        })

        q["attempted"] = True if attempt else False

    return jsonify(quizzes)
# ================= CHECK ATTEMPT =================
@app.route("/check_attempt/<quiz_id>")
def check_attempt(quiz_id):

    student_id = session.get("user")

    existing = submissions.find_one({
        "quiz_id": quiz_id,
        "student_id": student_id
    })

    return jsonify({"attempted": True if existing else False})

# ================= LOAD QUIZ PAGE =================
@app.route("/quiz/<quiz_id>")
def attend_quiz(quiz_id):
    return render_template("student_quiz.html")

# ================= GET QUIZ =================
@app.route("/get_quiz/<quiz_id>")
def get_quiz(quiz_id):
    q = quiz.find_one({"quiz_id":quiz_id},{"_id":0})
    return jsonify(q)

# ================= GET QUESTIONS =================
@app.route("/get_questions/<quiz_id>")
def get_questions(quiz_id):

    q=list(questions.find({"quiz_id":quiz_id},{"_id":0}))
    return jsonify(q)

# ================= SUBMIT QUIZ =================
@app.route("/submit_quiz",methods=["POST"])
def submit_quiz():

    data=request.json
    student_id = session.get("user")

    existing = submissions.find_one({
        "quiz_id": data["quiz_id"],
        "student_id": student_id
    })

    if existing:
        return jsonify({"msg":"Already submitted"})

    name = session.get("name")

    submissions.insert_one({
        "quiz_id": data["quiz_id"],
        "student_id": student_id,
        "name": name,
        "correct":data["correct"],
        "wrong":data["wrong"],
        "skipped":data["skipped"]
    })

    scores.insert_one({
        "quiz_id": data["quiz_id"],
        "student_id": student_id,
        "name": name,
        "correct":data["correct"],
        "wrong":data["wrong"],
        "skipped":data["skipped"],
        "result":"completed"
    })

    activity.insert_one({
        "quiz_id": data["quiz_id"],
        "student_id": student_id,
        "name": name,
        "question_answered": data["correct"] + data["wrong"],
        "correct": data["correct"],
        "wrong": data["wrong"],
        "skipped": data["skipped"],
        "violation_type": ", ".join([v["type"] for v in data.get("violations",[])]),
        "violation_count": len(data.get("violations",[])),
        "timestamp": datetime.now()
    })

    return jsonify({"msg":"submitted successfully"})

# ================= ACTIVITY =================
@app.route("/get_activity")
def get_activity():
    data=list(activity.find({},{"_id":0}))
    return jsonify(data)

# ================= SCORE =================
@app.route("/get_scores")
def get_scores():

    data=list(scores.find({},{"_id":0}))

    data=sorted(data,key=lambda x:x["correct"],reverse=True)

    for i,x in enumerate(data):
        if i==0: x["badge"]="🥇"
        elif i==1: x["badge"]="🥈"
        elif i==2: x["badge"]="🥉"
        else: x["badge"]="Bronze"

    return jsonify(data)



from twilio.twiml.messaging_response import MessagingResponse
import requests

whatsapp_sessions = db["whatsapp_sessions"]

@app.route("/whatsapp", methods=["POST"])
def whatsapp_webhook():

    resp = MessagingResponse()

    msg = request.form.get("Body")
    sender = request.form.get("From")

    media_url = request.form.get("MediaUrl0")
    media_type = request.form.get("MediaContentType0")

    # ================= GET USER SESSION =================
    user = whatsapp_sessions.find_one({"sender": sender})

    # ================= FILE UPLOAD =================
    if media_url:
        try:
            TWILIO_ACCOUNT_SID = os.getenv("TWILIO_ACCOUNT_SID")
            TWILIO_AUTH_TOKEN = os.getenv("TWILIO_AUTH_TOKEN")

            file_data = requests.get(
                media_url,
                auth=HTTPBasicAuth(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)
            ).content

            # FILE TYPE DETECTION
            if "pdf" in media_type:
                filename = "temp.pdf"
            elif "word" in media_type or "docx" in media_type:
                filename = "temp.docx"
            elif "text" in media_type:
                filename = "temp.txt"
            elif "csv" in media_type:
                filename = "temp.csv"
            else:
                resp.message("❌ Unsupported file type")
                return str(resp)

            with open(filename, "wb") as f:
                f.write(file_data)

            parsed_questions = []

            if filename.endswith(".pdf"):
                with pdfplumber.open(filename) as pdf:
                    lines = []
                    for p in pdf.pages:
                        t = p.extract_text()
                        if t:
                            lines += t.split("\n")
                parsed_questions = parse_block_questions(lines)

            elif filename.endswith(".docx"):
                doc = docx.Document(filename)
                lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                parsed_questions = parse_block_questions(lines)

            elif filename.endswith(".txt"):
                lines = file_data.decode().split("\n")
                parsed_questions = parse_block_questions(lines)

            elif filename.endswith(".csv"):
                df = pd.read_csv(filename)
                for _, r in df.iterrows():
                    parsed_questions.append({
                        "question": str(r["question"]),
                        "options": [r["A"], r["B"], r["C"], r["D"]],
                        "answer": str(r["answer"]).upper()
                    })

            if len(parsed_questions) == 0:
                resp.message("⚠️ No valid questions found in file")
                return str(resp)

            whatsapp_sessions.update_one(
                {"sender": sender},
                {"$set": {"questions": parsed_questions}},
                upsert=True
            )

            resp.message(f"✅ {len(parsed_questions)} questions uploaded successfully")
            return str(resp)

        except Exception as e:
            print("ERROR:", str(e))
            resp.message(f"❌ Error: {str(e)}")
            return str(resp)

    # ================= TEXT FLOW =================
    if not msg:
        resp.message("Send message")
        return str(resp)
    msg_lower = msg.lower()

    # STEP 1
    if "create quiz" in msg_lower:
        whatsapp_sessions.update_one(
            {"sender": sender},
            {"$set": {"step": "title", "data": {}}},
            upsert=True
        )
        resp.message("📘 Enter Quiz Title")
        return str(resp)

    # STEP 2: TITLE
    if user and user.get("step") == "title":
        whatsapp_sessions.update_one(
            {"sender": sender},
            {"$set": {"step": "duration", "data.title": msg}}
        )
        resp.message("⏱ Enter Duration (minutes)")
        return str(resp)

    # STEP 3: DURATION ✅ FIXED FOREVER
    if user and user.get("step") == "duration":
        try:
            duration = int(msg)

            whatsapp_sessions.update_one(
                {"sender": sender},
                {"$set": {"step": "start", "data.duration": duration}}
            )

            resp.message("📅 Enter Start Time (YYYY-MM-DD HH:MM)")
            return str(resp)

        except:
            resp.message("❌ Enter valid number like 20")
            return str(resp)

    # STEP 4: START TIME
    if user and user.get("step") == "start":
        try:
            start = datetime.strptime(msg, "%Y-%m-%d %H:%M")
            duration = user["data"]["duration"]

            end = start + timedelta(minutes=duration)
            quiz_id = str(uuid.uuid4())[:8]

            # SAVE QUIZ
            quiz.insert_one({
                "quiz_id": quiz_id,
                "title": user["data"]["title"],
                "start_time": start.isoformat(),
                "end_time": end.isoformat(),
                "duration": duration
            })

            # SAVE QUESTIONS
            for q in user.get("questions", []):
                questions.insert_one({
                    "quiz_id": quiz_id,
                    "question": q["question"],
                    "options": q["options"],
                    "answer": q["answer"]
                })

            # CLEAR SESSION
            whatsapp_sessions.delete_one({"sender": sender})

            resp.message(
                f"✅ Quiz Created!\n\n🔗 {request.host_url}join/{quiz_id}"
            )
            return str(resp)

        except:
            resp.message("❌ Format: 2026-04-01 10:30")
            return str(resp)

    resp.message("Say 'create quiz'")
    return str(resp)

# ================= RUN =================
if __name__ == "__main__":
    app.run(debug=True)
